# irt.py
# -*- coding: utf-8 -*-
"""
هسته‌ی IRT (مدل سه‌پارامتری 3PL) با پیاده‌سازی پایدار:

- احتمال 3PL با کلیپ و جلوگیری از overflow
- گرادیان صحیح لگ‌لایکلیهود 3PL نسبت به θ
- تخمین θ با روش نیوتن/فیشِر (پایدارتر از گام ثابت)
- تابع اطلاعات آیتم و آزمون (فرمول استاندارد 3PL)
- انتخاب سؤال بر مبنای بیشترین اطلاعات
- تخمین خطای استاندارد θ براساس اطلاعات آزمون
- ترسیم ICC و تابع اطلاعات
- خروجی نتایج به Excel و Word

نکته‌ی مهم:
منطق عددی و رفتاری توابع اصلی (three_pl_probability، item_information،
estimate_theta_mle/MAP/EAP، select_next_question) نسبت به نسخه‌ی قبلی
تغییر اساسی نکرده تا آزمون انطباقی موجود در پروژه خراب نشود.
"""

import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document

# ------------------------ ثوابت عددی ------------------------

EPS = 1e-9
THETA_MIN, THETA_MAX = -4.0, 4.0


# ------------------------ توابع پایه ------------------------

def _sigmoid(x: float) -> float:
    """
    تابع لاجیستیک با کلیپ کردن ورودی برای جلوگیری از overflow در exp.
    """
    x = np.clip(x, -50.0, 50.0)
    return 1.0 / (1.0 + np.exp(-x))


def three_pl_probability(theta: float, a: float, b: float, c: float) -> float:
    """
    احتمال پاسخ صحیح در مدل سه‌پارامتری (3PL):

        P(X=1 | θ) = c + (1 - c) * logistic(a(θ - b))

    - a: ضریب شیب (تمایز)
    - b: دشواری
    - c: پارامتر حد حدس

    برای پایداری عددی:
    - a حداقل EPS در نظر گرفته می‌شود.
    - c در بازه‌ی [0, 0.999] کلیپ می‌شود.
    - خروجی p در بازه‌ی [c+EPS, 1-EPS] کلیپ می‌شود تا log(0) و تقسیم‌های خطرناک رخ ندهد.
    """
    a = max(float(a), EPS)
    c = float(np.clip(c, 0.0, 0.999))

    p = c + (1.0 - c) * _sigmoid(a * (theta - b))
    # کلیپ نهایی احتمال برای جلوگیری از log(0) و ... در توابع بعدی
    return float(np.clip(p, c + EPS, 1.0 - EPS))


# ------------------------ درستنمایی و گرادیان ------------------------

def log_likelihood(theta: float, responses, item_params) -> float:
    """
    لگ‌لایکلیهود پاسخ‌ها در θ.

    پارامترها:
    - responses: لیست/آرایه‌ی 0/1 (درست/نادرست)
    - item_params: iterable از (a, b, c) برای هر آیتم
    """
    ll = 0.0
    for r, (a, b, c) in zip(responses, item_params):
        p = three_pl_probability(theta, a, b, c)
        ll += r * np.log(p) + (1 - r) * np.log(1.0 - p)
    return float(ll)


def _grad_loglik_theta(theta: float, responses, item_params) -> float:
    """
    گرادیان لگ‌لایکلیهود نسبت به θ در مدل 3PL:

        ∂/∂θ log L(θ) = Σ a_i * (x_i - p_i) * (p_i - c_i) / ((1 - c_i) * p_i)

    که در آن:
    - x_i پاسخ (0/1)
    - p_i = P(X_i=1 | θ)
    """
    g = 0.0
    for x, (a, b, c) in zip(responses, item_params):
        p = three_pl_probability(theta, a, b, c)
        g += a * (x - p) * (p - c) / ((1.0 - c) * p + EPS)
    return float(g)


# ------------------------ اطلاعات آیتم / آزمون و SE ------------------------

def item_information(theta: float, a: float, b: float, c: float) -> float:
    """
    اطلاعات فیشر آیتم در θ برای مدل 3PL:

        I_i(θ) = a^2 * ((p - c)^2 / (1 - c)^2) * (q / p)

    که p = P(X=1|θ) و q = 1 - p است.
    """
    p = three_pl_probability(theta, a, b, c)
    q = 1.0 - p
    denom = (1.0 - c) ** 2 + EPS
    info = (a ** 2) * ((p - c) ** 2) / denom * (q / p)
    return float(info) if np.isfinite(info) and info > 0 else 0.0


def test_information(theta: float, item_params) -> float:
    """
    اطلاعات کل آزمون در θ: جمع اطلاعات آیتم‌ها.
    """
    return float(sum(item_information(theta, a, b, c) for (a, b, c) in item_params))


def theta_se(theta: float, item_params) -> float:
    """
    خطای استاندارد تخمین θ بر اساس اطلاعات آزمون:

        SE(θ) ≈ 1 / sqrt(I_test(θ))

    برای جلوگیری از division by zero، I_test حداقل EPS در نظر گرفته می‌شود.
    """
    I = test_information(theta, item_params)
    return float(1.0 / np.sqrt(max(I, EPS)))


# ------------------------ تخمین θ (MLE / MAP / EAP) ------------------------

def estimate_theta_mle(responses, item_params,
                       lr: float = 0.01,
                       max_iter: int = 50,
                       tol: float = 1e-4) -> float:
    """
    تخمین θ به روش MLE با الگوریتم شبه نیوتن (Fisher scoring).

    پارامترها:
    - responses: لیست 0/1
    - item_params: iterable از (a, b, c)
    - lr: برای سازگاری امضای تابع حفظ شده، اما در عمل استفاده نمی‌شود
          (برای اینکه رفتار نسخه‌ی قبلی شما تغییری نکند).
    - max_iter: حداکثر تکرار
    - tol: آستانه‌ی همگرایی بر اساس |θ_{new} - θ_{old}|

    نکته:
    - step = g / I  (گرادیان تقسیم بر اطلاعات فیشر)
    - اگر step خیلی بزرگ/ناپایدار باشد، با tanh نرم می‌شود تا از پرش‌های
      عددی شدید جلوگیری شود.
    """
    theta = 0.0

    for _ in range(max_iter):
        g = _grad_loglik_theta(theta, responses, item_params)
        I = test_information(theta, item_params) + EPS
        step = g / I

        # نرم‌سازی گام‌های بسیار بزرگ یا ناپایدار
        if not np.isfinite(step) or abs(step) > 1.0:
            step = 0.25 * np.tanh(step)

        theta_new = float(np.clip(theta + step, THETA_MIN, THETA_MAX))
        if abs(theta_new - theta) < tol:
            return theta_new

        theta = theta_new

    return theta


def estimate_theta_map(responses, item_params,
                       theta0: float = 0.0,
                       prior_mean: float = 0.0,
                       prior_var: float = 1.0,
                       max_iter: int = 50,
                       tol: float = 1e-4) -> float:
    """
    تخمین θ به روش MAP با prior نرمال N(prior_mean, prior_var).

    - برای تعداد آیتم کم، MAP پایدارتر از MLE است.
    """
    theta = float(np.clip(theta0, THETA_MIN, THETA_MAX))
    inv_var = 1.0 / max(prior_var, EPS)

    for _ in range(max_iter):
        g_like = _grad_loglik_theta(theta, responses, item_params)
        g_prior = -(theta - prior_mean) * inv_var
        g = g_like + g_prior

        I = test_information(theta, item_params) + inv_var + EPS
        step = g / I

        if not np.isfinite(step) or abs(step) > 1.0:
            step = 0.25 * np.tanh(step)

        theta_new = float(np.clip(theta + step, THETA_MIN, THETA_MAX))
        if abs(theta_new - theta) < tol:
            return theta_new

        theta = theta_new

    return theta


def estimate_theta_eap(responses, item_params,
                       nodes: int = 41,
                       prior_mean: float = 0.0,
                       prior_sd: float = 1.0) -> float:
    """
    تخمین EAP (Expected A Posteriori) با تقریب گوس–هرمیت.

    - برای پاسخ‌های بسیار کم (مثلاً 1-2 آیتم) معمولاً پایدارتر از MLE است.
    - prior ~ N(prior_mean, prior_sd^2)
    """
    from numpy.polynomial.hermite import hermgauss

    # گره‌ها و وزن‌های گوس–هرمیت
    z, w = hermgauss(nodes)

    # نگاشت به N(prior_mean, prior_sd^2)
    thetas = np.sqrt(2.0) * prior_sd * z + prior_mean
    weights = w / np.sqrt(np.pi)  # وزن‌های prior نرمال

    def person_ll(th):
        s = 0.0
        for x, (a, b, c) in zip(responses, item_params):
            p = three_pl_probability(th, a, b, c)
            s += np.log(p if x else (1.0 - p))
        return s

    # لگ‌لایکلیهود روی شبکه‌ی θ
    logliks = np.array([person_ll(th) for th in thetas])
    m = np.max(logliks)

    # posterior ∝ exp(loglik) * prior  (که prior در weights است)
    post = np.exp(logliks - m) * weights
    post /= (np.sum(post) + EPS)

    eap = float(np.sum(thetas * post))
    return float(np.clip(eap, THETA_MIN, THETA_MAX))


# ------------------------ انتخاب سؤال ------------------------

def select_next_question(theta: float, all_item_params, answered_indices):
    """
    انتخاب آیتم با بیشترین اطلاعات در θ.

    پارامترها:
    - theta: مقدار فعلی θ
    - all_item_params: لیست/آرایه از (a, b, c) برای *همه‌ی* آیتم‌ها
    - answered_indices: مجموعه/لیست اندیس‌های از قبل پاسخ‌داده‌شده (0-مبنا)

    خروجی:
    - best_idx: اندیس آیتم انتخابی (0-مبنا) یا None اگر همه پاسخ داده شده باشند.
    """
    # برای کارایی بهتر، اگر لیست بود آن را به set تبدیل می‌کنیم (O(1) membership)
    answered_set = set(answered_indices) if answered_indices is not None else set()

    best_idx, best_info = None, -1.0
    for i, (a, b, c) in enumerate(all_item_params):
        if i in answered_set:
            continue
        info = item_information(theta, a, b, c)
        if info > best_info:
            best_info, best_idx = info, i

    return best_idx


# ------------------------ ترسیم نمودارها ------------------------

def plot_icc(item_params, save_path: str = 'static/icc.png') -> str:
    """
    ترسیم منحنی‌های ICC برای آیتم‌های داده‌شده.

    پارامتر:
    - item_params: iterable از (a, b, c) برای آیتم‌ها
    - save_path: مسیر ذخیره‌ی تصویر

    خروجی: همان save_path
    """
    theta_range = np.linspace(THETA_MIN, THETA_MAX, 200)

    plt.figure(figsize=(10, 6))
    for i, (a, b, c) in enumerate(item_params):
        probs = [three_pl_probability(t, a, b, c) for t in theta_range]
        plt.plot(theta_range, probs, label=f"سوال {i+1}")

    plt.xlabel('θ (توانایی)')
    plt.ylabel('احتمال پاسخ صحیح')
    plt.title('تابع مشخصه سوالات (ICC)')
    plt.legend(fontsize=8, ncol=2)
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(save_path)
    plt.close()

    return save_path


def plot_item_information(item_params, save_path: str = 'static/item_info.png') -> str:
    """
    ترسیم تابع اطلاعات کل آزمون در بازه‌ی θ.
    """
    theta_range = np.linspace(THETA_MIN, THETA_MAX, 200)
    total_info = np.zeros_like(theta_range, dtype=float)

    for a, b, c in item_params:
        info = np.array([item_information(t, a, b, c) for t in theta_range], dtype=float)
        total_info += info

    plt.figure(figsize=(8, 5))
    plt.plot(theta_range, total_info)
    plt.xlabel('θ (توانایی)')
    plt.ylabel('اطلاعات آزمون')
    plt.title('تابع اطلاعات کل آزمون')
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(save_path)
    plt.close()

    return save_path


# ------------------------ خروجی گرفتن ------------------------

def save_results_to_excel(filepath: str, responses, item_params, theta: float) -> None:
    """
    ذخیره‌ی نتایج در فایل Excel.

    فرض:
    - item_params فقط شامل آیتم‌های پاسخ‌داده‌شده است (هم‌اندازه با responses).
    - اگر می‌خواهی متن سؤال را هم اضافه کنی، می‌توانی قبل از صدا زدن این تابع
      یک ستون دیگر به DataFrame اضافه کنی یا wrapper بسازی.
    """
    df = pd.DataFrame({
        'سوال': [f"سوال {i + 1}" for i in range(len(responses))],
        'پاسخ (0/1)': responses,
        'a': [a for a, _, _ in item_params],
        'b': [b for _, b, _ in item_params],
        'c': [c for _, _, c in item_params],
    })
    df.loc[len(df.index)] = ['θ (توانایی)', theta, '', '', '']
    df.to_excel(filepath, index=False)


def save_results_to_word(filepath: str, responses, item_params, theta: float) -> None:
    """
    ذخیره‌ی نتایج در فایل Word (.docx).

    فرض:
    - item_params فقط شامل آیتم‌های پاسخ‌داده‌شده است (هم‌اندازه با responses).
    """
    doc = Document()
    doc.add_heading('نتایج آزمون انطباقی (3PL)', level=1)
    doc.add_paragraph(f'مقدار تخمینی θ: {theta:.3f}')

    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = 'سوال'
    hdr[1].text = 'پاسخ (0/1)'
    hdr[2].text = 'a'
    hdr[3].text = 'b'
    hdr[4].text = 'c'

    for i, (r, (a, b, c)) in enumerate(zip(responses, item_params)):
        row = table.add_row().cells
        row[0].text = f"سوال {i + 1}"
        row[1].text = str(int(r))
        row[2].text = f"{a:.3f}"
        row[3].text = f"{b:.3f}"
        row[4].text = f"{c:.3f}"

    doc.save(filepath)
