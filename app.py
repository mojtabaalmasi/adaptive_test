# app.py (Thesis-aligned + schema-robust + no column-name crashes)
from flask import (
    Flask, flash, render_template, request, redirect, session,
    send_file, url_for, send_from_directory, jsonify
)
import sqlite3
import numpy as np
import pandas as pd
from docx import Document
import matplotlib.pyplot as plt
import os
import uuid
from werkzeug.utils import secure_filename

# ----------------------------- پیکربندی پایه -----------------------------
app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret")

BASE_DIR = os.path.dirname(__file__)
DATABASE = os.environ.get("DATABASE_PATH", os.path.join(BASE_DIR, "questions.db"))

# ----------------------------- DB helpers -----------------------------
def get_db_connection():
    conn = sqlite3.connect(DATABASE, timeout=30)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON;")
    conn.execute("PRAGMA journal_mode=WAL;")
    conn.execute("PRAGMA synchronous=NORMAL;")
    conn.execute("PRAGMA busy_timeout=30000;")
    return conn

def table_columns(table_name: str) -> set:
    with get_db_connection() as conn:
        rows = conn.execute(f"PRAGMA table_info({table_name})").fetchall()
    return {r["name"] for r in rows}

def pick_text_column(table_name: str) -> str:
    """
    برای سازگاری با DBهای مختلف:
    اگر ستون text وجود داشت همان؛ وگرنه question_text؛
    اگر هیچ‌کدام نبود خطا می‌دهد تا شما بفهمید schema واقعاً چیست.
    """
    cols = table_columns(table_name)
    if "text" in cols:
        return "text"
    if "question_text" in cols:
        return "question_text"
    raise RuntimeError(f"Neither 'text' nor 'question_text' exists in table '{table_name}'. Columns={cols}")

# ----------------------------- Upload (Voice) -----------------------------
VOICE_BASE = os.environ.get("VOICE_PATH", os.path.join(BASE_DIR, "voices"))
os.makedirs(VOICE_BASE, exist_ok=True)

app.config["UPLOAD_FOLDER"] = VOICE_BASE
app.config["MAX_CONTENT_LENGTH"] = 15 * 1024 * 1024  # 15MB

MIME_EXT = {
    "audio/webm": "webm",
    "video/webm": "webm",
    "audio/ogg":  "ogg",
    "audio/mp4":  "m4a",
    "audio/mpeg": "mp3",
    "audio/3gpp": "3gp",
    "audio/wav":  "wav",
}

# فلگ سازگاری: آیا جدول answers ستونی به نام response دارد؟ (DBهای قدیمی)
_ANSWERS_HAS_RESPONSE_COL = None
def answers_has_response_column():
    global _ANSWERS_HAS_RESPONSE_COL
    if _ANSWERS_HAS_RESPONSE_COL is not None:
        return _ANSWERS_HAS_RESPONSE_COL
    cols = table_columns("answers")
    _ANSWERS_HAS_RESPONSE_COL = ("response" in cols)
    return _ANSWERS_HAS_RESPONSE_COL

# ----------------------------- شِمای جداول (ایجاد در صورت نبود) -----------------------------
def init_db():
    with sqlite3.connect(DATABASE, timeout=30) as conn:
        conn.execute("PRAGMA foreign_keys = ON;")
        conn.execute("PRAGMA busy_timeout=30000;")
        cur = conn.cursor()

        cur.executescript("""
        CREATE TABLE IF NOT EXISTS participants (
            participant_id    INTEGER PRIMARY KEY AUTOINCREMENT,
            name              TEXT,
            nationality       TEXT,
            mother_tongue     TEXT,
            official_language TEXT,
            age               INTEGER,
            major             TEXT,
            education_level   TEXT,
            job               TEXT,
            role              TEXT,
            current_level     TEXT,
            class_code        TEXT
        );

        CREATE TABLE IF NOT EXISTS teacher_info (
            participant_id               INTEGER UNIQUE,
            institution        TEXT,
            teaching_level               TEXT,
            importance_of_academic_persian TEXT,
            teaching_years    INTEGER,
            FOREIGN KEY(participant_id) REFERENCES participants(participant_id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS learner_info (
            participant_id               INTEGER UNIQUE,
            learning_duration            TEXT,
            current_level                TEXT,
            formal_training              TEXT,
            training_institution         TEXT,
            samfa_taken                  TEXT,
            samfa_score                  REAL,
            importance_of_academic_persian TEXT,
            speaking_ability             TEXT,
            reading_ability              TEXT,
            writing_ability              TEXT,
            listening_ability            TEXT,
            FOREIGN KEY(participant_id) REFERENCES participants(participant_id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            text TEXT NOT NULL,
            option1 TEXT, option2 TEXT, option3 TEXT, option4 TEXT,
            correct_option INTEGER,
            a REAL NOT NULL, b REAL NOT NULL, c REAL NOT NULL
        );

        CREATE TABLE IF NOT EXISTS answers (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id         INTEGER NOT NULL,
            question_id     INTEGER NOT NULL,
            selected_option INTEGER,
            is_correct      INTEGER,
            created_at      DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_id)     REFERENCES participants(participant_id) ON DELETE CASCADE,
            FOREIGN KEY(question_id) REFERENCES questions(id)                ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS user_results (
            user_id INTEGER PRIMARY KEY,
            theta   REAL,
            FOREIGN KEY(user_id) REFERENCES participants(participant_id) ON DELETE CASCADE
        );

        -- توجه: بعضی DBها ستون را question_text دارند؛ بعضی text.
        -- ما اینجا جدول را با text می‌سازیم، ولی کد با pick_text_column سازگار می‌شود.
        CREATE TABLE IF NOT EXISTS manager_questions (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            text          TEXT NOT NULL,
            display_order INTEGER NOT NULL DEFAULT 0,
            is_required   INTEGER NOT NULL DEFAULT 1 CHECK (is_required IN (0,1)),
            is_active     INTEGER NOT NULL DEFAULT 1 CHECK (is_active IN (0,1)),
            created_at    DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS voice_answers (
            id             INTEGER PRIMARY KEY AUTOINCREMENT,
            participant_id INTEGER NOT NULL,
            role           TEXT NOT NULL DEFAULT 'manager',
            question_id    INTEGER NOT NULL,
            file_path      TEXT NOT NULL,
            mime_type      TEXT,
            duration_ms    INTEGER,
            size_bytes     INTEGER,
            created_at     DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(participant_id) REFERENCES participants(participant_id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS strategies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            strategy TEXT NOT NULL,
            category TEXT,
            frequency INTEGER DEFAULT 0,
            target_role TEXT NOT NULL DEFAULT 'learner'
        );

        CREATE TABLE IF NOT EXISTS manager_info (
            participant_id INTEGER PRIMARY KEY,
            center_name TEXT,
            center_city TEXT,
            center_type TEXT,
            years_as_manager INTEGER,
            num_teachers INTEGER,
            num_learners INTEGER,
            FOREIGN KEY(participant_id) REFERENCES participants(participant_id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS teacher_post_questions (
            id INTEGER PRIMARY KEY,
            text TEXT NOT NULL,
            dimension TEXT,
            question_type TEXT DEFAULT 'open',
            scale TEXT,
            display_order INTEGER NOT NULL DEFAULT 0,
            is_active INTEGER NOT NULL DEFAULT 1,
            is_required INTEGER NOT NULL DEFAULT 1
        );

        CREATE TABLE IF NOT EXISTS manager_post_questions (
            id INTEGER PRIMARY KEY,
            text TEXT NOT NULL,
            dimension TEXT,
            question_type TEXT DEFAULT 'open',
            scale TEXT,
            display_order INTEGER NOT NULL DEFAULT 0,
            is_active INTEGER NOT NULL DEFAULT 1,
            is_required INTEGER NOT NULL DEFAULT 1
        );

        CREATE TABLE IF NOT EXISTS teacher_post_answers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            participant_id INTEGER NOT NULL,
            question_id INTEGER NOT NULL,
            answer_value INTEGER,
            answer_text TEXT,
            created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(participant_id) REFERENCES participants(participant_id) ON DELETE CASCADE,
            UNIQUE(participant_id, question_id)
        );

        CREATE TABLE IF NOT EXISTS manager_post_answers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            participant_id INTEGER NOT NULL,
            question_id INTEGER NOT NULL,
            answer_value INTEGER,
            answer_text TEXT,
            answered_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(participant_id) REFERENCES participants(participant_id) ON DELETE CASCADE,
            UNIQUE(participant_id, question_id)
        );

        CREATE TABLE IF NOT EXISTS test_sessions (
            session_id TEXT PRIMARY KEY,
            user_id INTEGER NOT NULL,
            role TEXT,
            started_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
            ended_at DATETIME,
            stop_reason TEXT,
            items_administered INTEGER DEFAULT 0,
            theta_start REAL,
            theta_final REAL,
            se_final REAL,
            FOREIGN KEY(user_id) REFERENCES participants(participant_id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS answers_meta (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id TEXT NOT NULL,
            participant_id INTEGER NOT NULL,
            question_id INTEGER NOT NULL,
            step INTEGER NOT NULL,
            selected_option INTEGER,
            is_correct INTEGER,
            theta_before REAL,
            theta_after REAL,
            se_after REAL,
            info REAL,
            created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(session_id) REFERENCES test_sessions(session_id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS strategy_answers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            participant_id INTEGER NOT NULL,
            strategy_id INTEGER NOT NULL,
            choice INTEGER NOT NULL,
            created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME,
            FOREIGN KEY(participant_id) REFERENCES participants(participant_id) ON DELETE CASCADE,
            FOREIGN KEY(strategy_id) REFERENCES strategies(id) ON DELETE CASCADE,
            UNIQUE(participant_id, strategy_id)
        );

        CREATE TABLE IF NOT EXISTS consents (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            participant_id  INTEGER NOT NULL,
            consented       INTEGER NOT NULL CHECK(consented IN (0,1)),
            consent_version TEXT,
            consented_at    DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS pkrq_items (
            item_id INTEGER PRIMARY KEY AUTOINCREMENT,
            text_fa TEXT NOT NULL,
            construct TEXT,
            construct_fa TEXT,
            is_active INTEGER DEFAULT 1,
            display_order INTEGER DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS pkrq_responses (
            response_id INTEGER PRIMARY KEY AUTOINCREMENT,
            participant_id INTEGER NOT NULL,
            item_id INTEGER NOT NULL,
            response_value INTEGER NOT NULL,
            mode TEXT NOT NULL CHECK(mode IN ('PBT','CBT','CAT')),
            answered_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(participant_id, item_id, mode)
        );

        -- ✅ اضافه شد: جدول نتایج نهایی آزمون
        CREATE TABLE IF NOT EXISTS test_results (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            participant_id INTEGER NOT NULL,
            mode TEXT NOT NULL,
            test_phase TEXT,
            raw_score REAL,
            theta REAL,
            se_final REAL,
            items_count INTEGER,
            duration_sec INTEGER,
            taken_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );
        """)

        # ایندکس/قیود مهم برای UPSERT ها
        cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS uq_voice_answers_user_q ON voice_answers(participant_id, question_id);")
        cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS uq_user_results_user ON user_results(user_id);")
        cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS uq_pkrq_response ON pkrq_responses(participant_id, item_id, mode);")

        conn.commit()

# ----------------------------- لایهٔ سؤال/پارامترها -----------------------------
def get_question_by_id(question_id: int):
    with get_db_connection() as conn:
        row = conn.execute(
            "SELECT id, text, option1, option2, option3, option4 FROM questions WHERE id = ?",
            (question_id,)
        ).fetchone()
    if row:
        return {"id": int(row["id"]), "text": row["text"],
                "options": [row["option1"], row["option2"], row["option3"], row["option4"]]}
    return None

def get_all_item_params():
    with get_db_connection() as conn:
        rows = conn.execute("SELECT id, a, b, c FROM questions ORDER BY id").fetchall()
    return [(int(r["id"]), float(r["a"]), float(r["b"]), float(r["c"])) for r in rows]

def get_correct_answer(question_id: int):
    with get_db_connection() as db:
        row = db.execute("SELECT correct_option FROM questions WHERE id = ?", (question_id,)).fetchone()
    return int(row["correct_option"]) if (row and row["correct_option"] is not None) else None

# ----------------------------- IRT (3PL پایدار) -----------------------------
EPS = 1e-9
THETA_MIN, THETA_MAX = -4.0, 4.0

def _sigmoid(x):
    x = np.clip(x, -50.0, 50.0)
    return 1.0 / (1.0 + np.exp(-x))

def three_pl_probability(theta, a, b, c):
    a = max(float(a), EPS)
    c = float(np.clip(c, 0.0, 0.999))
    p = c + (1.0 - c) * _sigmoid(a * (theta - b))
    return float(np.clip(p, c + EPS, 1.0 - EPS))

def item_information(theta, a, b, c):
    p = three_pl_probability(theta, a, b, c)
    q = 1.0 - p
    denom = (1.0 - c) ** 2 + EPS
    info = (a ** 2) * ((p - c) ** 2) / denom * (q / p)
    return float(info) if np.isfinite(info) and info > 0 else 0.0

def test_information(theta, item_params):
    return float(sum(item_information(theta, a, b, c) for (a, b, c) in item_params))

def theta_se(theta, item_params):
    I = test_information(theta, item_params)
    return float(1.0 / np.sqrt(max(I, EPS)))

def _grad_loglik_theta(theta, responses, item_params):
    g = 0.0
    for x, (a, b, c) in zip(responses, item_params):
        p = three_pl_probability(theta, a, b, c)
        g += a * (x - p) * (p - c) / ((1.0 - c) * p + EPS)
    return float(g)

def estimate_theta_mle(responses, item_params, max_iter=50, tol=1e-4):
    theta = 0.0
    for _ in range(max_iter):
        g = _grad_loglik_theta(theta, responses, item_params)
        I = test_information(theta, item_params) + EPS
        step = g / I
        if not np.isfinite(step) or abs(step) > 1.0:
            step = 0.25 * np.tanh(step)
        theta_new = float(np.clip(theta + step, THETA_MIN, THETA_MAX))
        if abs(theta_new - theta) < tol:
            return theta_new
        theta = theta_new
    return theta

def estimate_theta_map(responses, item_params, theta0=0.0, prior_mean=0.0, prior_var=1.0, max_iter=50, tol=1e-4):
    theta = float(np.clip(theta0, THETA_MIN, THETA_MAX))
    inv_var = 1.0 / max(prior_var, EPS)
    for _ in range(max_iter):
        g_like = _grad_loglik_theta(theta, responses, item_params)
        g_prior = -(theta - prior_mean) * inv_var
        g = g_like + g_prior
        I = test_information(theta, item_params) + inv_var + EPS
        step = g / I
        if not np.isfinite(step) or abs(step) > 1.0:
            step = 0.25 * np.tanh(step)
        theta_new = float(np.clip(theta + step, THETA_MIN, THETA_MAX))
        if abs(theta_new - theta) < tol:
            return theta_new
        theta = theta_new
    return theta

def select_next_question(theta, all_item_params, answered_indices):
    best_idx, best_info = None, -1.0
    for i, (a, b, c) in enumerate(all_item_params):
        if i in answered_indices:
            continue
        info = item_information(theta, a, b, c)
        if info > best_info:
            best_info, best_idx = info, i
    return best_idx

# ----------------------------- نمودار -----------------------------
def plot_icc(item_params, save_path):
    theta_range = np.linspace(THETA_MIN, THETA_MAX, 200)
    plt.figure(figsize=(10, 6))
    for i, (a, b, c) in enumerate(item_params):
        probs = [three_pl_probability(t, a, b, c) for t in theta_range]
        plt.plot(theta_range, probs, label=f"سوال {i+1}")
    plt.xlabel("θ (توانایی)")
    plt.ylabel("احتمال پاسخ صحیح")
    plt.title("تابع مشخصه سوالات (ICC)")
    plt.legend(fontsize=8, ncol=2)
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(save_path)
    plt.close()
    return save_path

def plot_item_information(item_params, save_path):
    theta_range = np.linspace(THETA_MIN, THETA_MAX, 200)
    total_info = np.zeros_like(theta_range, dtype=float)
    for a, b, c in item_params:
        info = np.array([item_information(t, a, b, c) for t in theta_range], dtype=float)
        total_info += info
    plt.figure(figsize=(8, 5))
    plt.plot(theta_range, total_info)
    plt.xlabel("θ (توانایی)")
    plt.ylabel("اطلاعات آزمون")
    plt.title("تابع اطلاعات کل آزمون")
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(save_path)
    plt.close()
    return save_path

# ----------------------------- خروجی به Excel/Word -----------------------------
def save_results_to_excel(filepath, responses, answered_indices, theta):
    with get_db_connection() as conn:
        rows = conn.execute("SELECT text, a, b, c FROM questions ORDER BY id").fetchall()
    data = []
    for i, r in zip(answered_indices, responses):
        row = rows[i]
        data.append({"سوال": row["text"], "پاسخ (0/1)": r, "a": row["a"], "b": row["b"], "c": row["c"]})
    df = pd.DataFrame(data)
    df.loc[len(df)] = ["θ (توانایی)", theta, "", "", ""]
    df.to_excel(filepath, index=False)

def save_results_to_word(filepath, responses, answered_indices, theta):
    with get_db_connection() as conn:
        rows = conn.execute("SELECT text, a, b, c FROM questions ORDER BY id").fetchall()
    doc = Document()
    doc.add_heading("نتایج آزمون تطبیقی (3PL)", 0)
    doc.add_paragraph(f"مقدار تخمینی θ: {theta:.3f}")
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "سوال", "پاسخ (0/1)", "a", "b", "c"
    for i, r in zip(answered_indices, responses):
        row = rows[i]
        cells = table.add_row().cells
        cells[0].text = row["text"]
        cells[1].text = str(r)
        cells[2].text = str(row["a"])
        cells[3].text = str(row["b"])
        cells[4].text = str(row["c"])
    doc.save(filepath)

# ----------------------------- مسیرهای وب -----------------------------
@app.route("/")
def index():
    session.clear()
    return render_template("index.html")

@app.route("/voices/<path:filepath>")
def serve_voice(filepath):
    return send_from_directory(VOICE_BASE, filepath, as_attachment=False)

# ----------------------------- رضایت آگاهانه -----------------------------
@app.route("/consent", methods=["GET", "POST"])
def consent():
    if request.method == "POST":
        consented = 1 if request.form.get("consented") in ("1", "on", "true", "True") else 0
        session["consented"] = consented
        session["consent_version"] = request.form.get("consent_version", "v1")
        if consented != 1:
            flash("برای ادامه، لازم است با فرم رضایت آگاهانه موافقت کنید.", "warning")
            return redirect(url_for("consent"))
        return redirect(url_for("register"))
    return render_template("consent.html")

# ----------------------------- پرسشنامه نگرش (PKRQ) -----------------------------
@app.route("/attitude/<mode>", methods=["GET", "POST"])
def attitude(mode):
    if not session.get("pending_pkrq"):
        return redirect(url_for("test"))

    mode = (mode or "").upper()
    if mode not in ("PBT", "CBT", "CAT"):
        return "Mode نامعتبر است.", 400
    if "participant_id" not in session:
        return redirect(url_for("index"))

    with get_db_connection() as conn:
        items = conn.execute(
            "SELECT item_id AS id, text_fa AS text FROM pkrq_items "
            "WHERE COALESCE(is_active,1)=1 "
            "ORDER BY COALESCE(display_order,0), item_id"
        ).fetchall()

    if not items:
        flash("پرسشنامه نگرش در پایگاه داده تعریف نشده است.", "warning")
        return redirect(url_for("thank_you"))

    if request.method == "POST":
        pid = int(session["participant_id"])
        with get_db_connection() as conn:
            cur = conn.cursor()
            for it in items:
                key = f"item_{it['id']}"
                val = request.form.get(key)
                if val is None:
                    continue
                try:
                    v = int(val)
                except ValueError:
                    continue
                cur.execute(
                    """
                    INSERT INTO pkrq_responses (participant_id, item_id, response_value, mode)
                    VALUES (?, ?, ?, ?)
                    ON CONFLICT(participant_id, item_id, mode)
                    DO UPDATE SET response_value=excluded.response_value, answered_at=CURRENT_TIMESTAMP
                    """,
                    (pid, int(it["id"]), v, mode),
                )
            conn.commit()

        flash("پاسخ‌های نگرش ثبت شد.", "success")
        session["pending_pkrq"] = False
        return redirect(url_for("post_test"))

    return render_template("pkrq.html", items=items, mode=mode)

# ----------------------------- پرسشنامه/مصاحبه مدیر + صوت -----------------------------
@app.route("/manager_survey", methods=["GET"])
def manager_survey():
    if "participant_id" not in session:
        return redirect(url_for("index"))
    if session.get("role") != "manager":
        return redirect(url_for("test"))

    txt_col = pick_text_column("manager_questions")
    with get_db_connection() as conn:
        questions = conn.execute(
            f"SELECT id, {txt_col} AS text FROM manager_questions "
            f"WHERE COALESCE(is_active,1)=1 ORDER BY COALESCE(display_order,0), id"
        ).fetchall()
    return render_template("manager_survey.html", questions=questions)

@app.route("/api/voice_answer", methods=["POST"])
def api_voice_answer():
    if "participant_id" not in session:
        return jsonify(ok=False, error="not_authenticated"), 401

    pid = int(session["participant_id"])
    qid = request.form.get("question_id")
    if not qid:
        return jsonify(ok=False, error="missing_question_id"), 400
    try:
        qid_int = int(qid)
    except ValueError:
        return jsonify(ok=False, error="bad_question_id"), 400

    f = request.files.get("audio")
    if not f or f.filename == "":
        return jsonify(ok=False, error="missing_audio"), 400

    mime = (f.mimetype or "").lower()
    ext = MIME_EXT.get(mime) or (os.path.splitext(f.filename)[1].lstrip(".") or "bin")

    filename = secure_filename(f"{pid}_{qid_int}_{uuid.uuid4().hex}.{ext}")
    abs_path = os.path.join(VOICE_BASE, filename)
    f.save(abs_path)

    rel_path = filename  # served as /voices/<filename>

    with get_db_connection() as conn:
        conn.execute(
            """
            INSERT INTO voice_answers (participant_id, question_id, file_path, mime_type, size_bytes)
            VALUES (?, ?, ?, ?, ?)
            ON CONFLICT(participant_id, question_id)
            DO UPDATE SET file_path=excluded.file_path, mime_type=excluded.mime_type, size_bytes=excluded.size_bytes, created_at=CURRENT_TIMESTAMP
            """,
            (pid, qid_int, rel_path, mime or None, os.path.getsize(abs_path)),
        )
        conn.commit()

    return jsonify(ok=True, file_path=rel_path)

# ----------------------------- ثبت‌نام -----------------------------
@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        try:
            with sqlite3.connect(DATABASE, timeout=30) as conn:
                conn.row_factory = sqlite3.Row
                conn.execute("PRAGMA foreign_keys = ON;")
                conn.execute("PRAGMA busy_timeout=30000;")
                cur = conn.cursor()

                name = request.form.get("name")
                nationality = request.form.get("nationality")
                mother_tongue = request.form.get("mother_tongue")
                official_language = request.form.get("official_language")
                age = request.form.get("age")
                major = request.form.get("major")
                education_level = request.form.get("education_level")
                job = request.form.get("job")
                role = request.form.get("role")

                cur.execute(
                    """
                    INSERT INTO participants
                    (name, nationality, mother_tongue, official_language, age, major, education_level, job, role)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (name, nationality, mother_tongue, official_language, age, major, education_level, job, role),
                )
                participant_id = cur.lastrowid

                if role == "teacher":
                    cur.execute(
                        """
                        INSERT INTO teacher_info
                        (participant_id, teaching_years, institution, teaching_level, importance_of_academic_persian)
                        VALUES (?, ?, ?, ?, ?)
                        """,
                        (
                            participant_id,
                            request.form.get("teaching_years"),
                            request.form.get("institution"),
                            request.form.get("teaching_level"),
                            request.form.get("academic_persian_opinion"),
                        ),
                    )

                elif role == "learner":
                    cur.execute(
                        """
                        INSERT INTO learner_info
                        (participant_id, learning_duration, current_level, formal_training,
                         training_institution, samfa_taken, samfa_score, importance_of_academic_persian,
                         speaking_ability, reading_ability, writing_ability, listening_ability)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            participant_id,
                            request.form.get("learning_duration"),
                            request.form.get("current_level"),
                            request.form.get("formal_training"),
                            request.form.get("institution"),
                            request.form.get("samfa_taken"),
                            request.form.get("samfa_score"),
                            request.form.get("importance_of_academic_persian"),
                            request.form.get("speaking_ability"),
                            request.form.get("reading_ability"),
                            request.form.get("writing_ability"),
                            request.form.get("listening_ability"),
                        ),
                    )

                elif role == "manager":
                    cur.execute(
                        """
                        INSERT INTO manager_info
                        (participant_id, center_name, center_city, center_type, years_as_manager, num_teachers, num_learners)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            participant_id,
                            request.form.get("center_name"),
                            request.form.get("center_city"),
                            request.form.get("center_type"),
                            request.form.get("years_as_manager"),
                            request.form.get("num_teachers"),
                            request.form.get("num_learners"),
                        ),
                    )

                conn.commit()

            session["participant_id"] = int(participant_id)
            session["user_name"] = name
            session["role"] = role

            # ثبت رضایت آگاهانه
            try:
                if session.get("consented") == 1:
                    with sqlite3.connect(DATABASE, timeout=30) as db:
                        db.execute("PRAGMA busy_timeout=30000;")
                        db.execute(
                            "INSERT INTO consents (participant_id, consented, consent_version) VALUES (?, ?, ?)",
                            (int(participant_id), 1, session.get("consent_version", "v1")),
                        )
                        db.commit()
            except Exception:
                pass

            if role == "manager":
                return redirect(url_for("manager_survey"))
            return redirect(url_for("test"))

        except sqlite3.Error as e:
            return f"خطای پایگاه داده: {e}", 500

    return render_template("register.html")

# ----------------------------- آزمون تطبیقی -----------------------------
@app.route("/test", methods=["GET", "POST"])
def test():
    if "participant_id" not in session:
        return redirect(url_for("index"))

    session.setdefault("answered_questions", [])
    session.setdefault("responses", [])
    session.setdefault("theta", 0.0)
    session.setdefault("stable_streak", 0)

    session.setdefault("test_session_id", None)
    if not session.get("test_session_id"):
        session["test_session_id"] = str(uuid.uuid4())
        with sqlite3.connect(DATABASE, timeout=30) as db:
            db.execute("PRAGMA busy_timeout=30000;")
            db.execute(
                "INSERT OR IGNORE INTO test_sessions (session_id, user_id, role, theta_start) VALUES (?, ?, ?, ?)",
                (session["test_session_id"], int(session["participant_id"]), session.get("role"), float(session.get("theta", 0.0))),
            )
            db.commit()

    answered = list(map(int, session["answered_questions"]))
    responses = list(map(int, session["responses"]))
    theta = float(session["theta"])
    streak = int(session["stable_streak"])

    rows = get_all_item_params()
    if not rows:
        flash("بانک سؤال خالی است.", "error")
        return redirect(url_for("index"))

    question_ids = [r[0] for r in rows]
    all_item_params = [tuple(r[1:]) for r in rows]
    total_questions = len(all_item_params)

    MIN_QUESTIONS = 8
    HARD_MAX = 22
    SE_TARGET = 0.30
    DELTA_TARGET = 0.03
    STREAK_NEED = 2
    CUT_SCORE = 0.0
    Z_CI = 1.96

    if request.method == "GET" and not answered:
        theta = 0.0
        responses = []
        streak = 0

        start_idx = select_next_question(theta, all_item_params, answered_indices=[])
        if start_idx is None:
            flash("سؤالی برای شروع یافت نشد.", "error")
            return redirect(url_for("index"))

        answered = [start_idx]
        session["answered_questions"] = answered
        session["responses"] = responses
        session["theta"] = float(theta)
        session["stable_streak"] = int(streak)

        current_qid = question_ids[answered[-1]]
        question = get_question_by_id(current_qid)
        progress = int(len(answered) / max(total_questions, 1) * 100)
        return render_template("test.html", question=question, progress=progress)

    if request.method == "POST":
        try:
            sel = int(request.form.get("answer"))
        except Exception:
            sel = None

        if sel not in (1, 2, 3, 4):
            current_qid = question_ids[answered[-1]]
            question = get_question_by_id(current_qid)
            progress = int(len(answered) / max(total_questions, 1) * 100)
            return render_template("test.html", question=question, error="گزینهٔ معتبر انتخاب نشده است.", progress=progress)

        current_idx = answered[-1]
        current_qid = question_ids[current_idx]

        co = get_correct_answer(current_qid)
        is_correct = 1 if (co is not None and sel == int(co)) else 0
        responses.append(is_correct)

        answered_params = [all_item_params[i] for i in answered]
        old_theta = theta
        theta = estimate_theta_map(responses, answered_params) if len(responses) < 3 else estimate_theta_mle(responses, answered_params)
        theta_change = abs(theta - old_theta)
        se_now = theta_se(theta, answered_params)

        participant_id = int(session["participant_id"])
        with sqlite3.connect(DATABASE, timeout=30) as db:
            db.execute("PRAGMA busy_timeout=30000;")
            cur = db.cursor()
            has_resp_col = answers_has_response_column()

            if has_resp_col:
                cur.execute(
                    "INSERT INTO answers (user_id, question_id, response, selected_option, is_correct) VALUES (?, ?, ?, ?, ?)",
                    (participant_id, current_qid, sel, sel, is_correct),
                )
            else:
                cur.execute(
                    "INSERT INTO answers (user_id, question_id, selected_option, is_correct) VALUES (?, ?, ?, ?)",
                    (participant_id, current_qid, sel, is_correct),
                )

            # answers_meta
            try:
                info_val = item_information(old_theta, *all_item_params[current_idx])
                cur.execute(
                    """
                    INSERT INTO answers_meta
                    (session_id, participant_id, question_id, step, selected_option, is_correct, theta_before, theta_after, se_after, info)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        session.get("test_session_id"),
                        participant_id,
                        current_qid,
                        len(responses),
                        sel,
                        is_correct,
                        float(old_theta),
                        float(theta),
                        float(se_now),
                        float(info_val),
                    ),
                )
            except sqlite3.Error:
                pass

            db.commit()

        streak = streak + 1 if theta_change < DELTA_TARGET else 0
        session["stable_streak"] = int(streak)

        num_answered = len(responses)
        above_cut = (theta - Z_CI * se_now) > CUT_SCORE
        below_cut = (theta + Z_CI * se_now) < CUT_SCORE

        stop_reason = None
        if num_answered >= MIN_QUESTIONS:
            if se_now <= SE_TARGET:
                stop_reason = f"دقت کافی (SE ≤ {SE_TARGET})"
            elif above_cut:
                stop_reason = f"نتیجهٔ قطعی: بالاتر از مرز {CUT_SCORE}"
            elif below_cut:
                stop_reason = f"نتیجهٔ قطعی: پایین‌تر از مرز {CUT_SCORE}"
            elif streak >= STREAK_NEED:
                stop_reason = f"پایداری θ (Δθ < {DELTA_TARGET} برای {STREAK_NEED} بار پیاپی)"

        if stop_reason is None and num_answered >= HARD_MAX:
            stop_reason = f"رسیدن به سقف {HARD_MAX} سؤال"

        if stop_reason is not None:
            with sqlite3.connect(DATABASE, timeout=30) as conn:
                conn.execute("PRAGMA busy_timeout=30000;")
                cur = conn.cursor()

                # ذخیره theta در user_results
                cur.execute(
                    "UPDATE user_results SET theta=? WHERE user_id=?",
                    (float(theta), participant_id),
                )
                if cur.rowcount == 0:
                    cur.execute(
                        "INSERT INTO user_results (user_id, theta) VALUES (?, ?)",
                        (participant_id, float(theta)),
                    )

                # بستن test_sessions
                try:
                    cur.execute(
                        """
                        UPDATE test_sessions
                        SET ended_at=CURRENT_TIMESTAMP,
                            stop_reason=?,
                            items_administered=?,
                            theta_final=?,
                            se_final=?
                        WHERE session_id=?
                        """,
                        (
                            stop_reason,
                            len(responses),
                            float(theta),
                            float(se_now),
                            session.get("test_session_id"),
                        ),
                    )
                except sqlite3.Error:
                    pass

                # ✅ درج قطعی در test_results (داخل همین تراکنش)
                try:
                    dur_row = cur.execute(
                        """
                        SELECT CAST((julianday('now') - julianday(started_at)) * 86400 AS INTEGER) AS dur
                        FROM test_sessions
                        WHERE session_id=?
                        """,
                        (session.get("test_session_id"),),
                    ).fetchone()
                    duration_sec = int(dur_row[0]) if dur_row and dur_row[0] is not None else None

                    cur.execute(
                        """
                        INSERT INTO test_results
                        (participant_id, mode, test_phase, raw_score, theta, se_final, items_count, duration_sec)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            participant_id,
                            "CAT",
                            None,
                            None,
                            float(theta),
                            float(se_now),
                            int(len(responses)),
                            duration_sec,
                        ),
                    )
                except sqlite3.Error as e:
                    # اینجا دیگر silent نباشد تا اگر mismatch بود بفهمیم
                    raise

            conn.commit()

            # فلگ‌ها و ریدایرکت بعد از commit
            session["theta"] = float(theta)
            session["answered_questions"] = answered
            session["responses"] = responses
            session["stop_reason"] = stop_reason
            session["pending_pkrq"] = True
            session["pending_post_test"] = True
            return redirect(url_for("attitude", mode="CAT"))

        next_idx = select_next_question(theta, all_item_params, answered)
        if next_idx is None:
            session["stop_reason"] = "پایان بانک سؤال"
            return redirect(url_for("result"))

        answered.append(next_idx)
        session["answered_questions"] = answered
        session["responses"] = responses
        session["theta"] = float(theta)

        next_qid = question_ids[next_idx]
        question = get_question_by_id(next_qid)
        progress = int(len(answered) / max(total_questions, 1) * 100)
        return render_template("test.html", question=question, progress=progress)

    current_qid = question_ids[answered[-1]]
    question = get_question_by_id(current_qid)
    progress = int(len(answered) / max(total_questions, 1) * 100)
    return render_template("test.html", question=question, progress=progress)

# ----------------------------- پس‌آزمون راهبردها (learner) -----------------------------
@app.route("/post_test", methods=["GET", "POST"])
def post_test():
    if not session.get("pending_post_test"):
        return redirect(url_for("result"))

    if "participant_id" not in session:
        return redirect(url_for("index"))

    participant_id = int(session["participant_id"])
    role = session.get("role", "learner")

    with sqlite3.connect(DATABASE, timeout=30) as conn:
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA busy_timeout=30000;")
        cur = conn.cursor()
        rows = cur.execute(
            """
            SELECT id, strategy AS text, COALESCE(category,'') AS category
            FROM strategies
            WHERE target_role=?
            ORDER BY category, id
            """,
            (role,),
        ).fetchall()

    if not rows:
        flash("هیچ موردی در پرسشنامهٔ راهبردها تعریف نشده است.", "error")
        return redirect(url_for("result"))

    groups = []
    last_cat = None
    for r in rows:
        cat = r["category"]
        if cat != last_cat:
            groups.append({"category": cat, "items": []})
            last_cat = cat
        groups[-1]["items"].append({"id": int(r["id"]), "text": r["text"]})

    if request.method == "POST":
        errors = {}
        payload = {}

        for g in groups:
            for item in g["items"]:
                key = f"s_{item['id']}"
                val = request.form.get(key)
                if not val:
                    errors[item["id"]] = "الزامی"
                    continue
                try:
                    choice = int(val)
                except ValueError:
                    errors[item["id"]] = "نامعتبر"
                    continue
                if not (1 <= choice <= 5):
                    errors[item["id"]] = "بازه ۱ تا ۵"
                    continue
                payload[item["id"]] = choice

        if errors:
            return render_template("strategies_survey.html", groups=groups, errors=errors, values=request.form)

        with sqlite3.connect(DATABASE, timeout=30) as conn:
            conn.execute("PRAGMA busy_timeout=30000;")
            cur = conn.cursor()
            for sid, choice in payload.items():
                cur.execute(
                    """
                    INSERT INTO strategy_answers (participant_id, strategy_id, choice)
                    VALUES (?, ?, ?)
                    ON CONFLICT(participant_id, strategy_id)
                    DO UPDATE SET choice=excluded.choice, updated_at=CURRENT_TIMESTAMP
                    """,
                    (participant_id, sid, choice),
                )
            conn.commit()

        session["pending_post_test"] = False
        return redirect(url_for("result"))

    return render_template("strategies_survey.html", groups=groups, errors={}, values={})

# ----------------------------- قدردانی -----------------------------
@app.route("/thank_you")
def thank_you():
    if "participant_id" not in session:
        return redirect(url_for("index"))
    return render_template("thank_you.html", user_name=session.get("user_name", "کاربر گرامی"))

# ----------------------------- نتیجه -----------------------------
@app.route("/result")
def result():
    if "responses" not in session or "answered_questions" not in session:
        return redirect(url_for("index"))

    responses = list(map(int, session["responses"]))
    answered = list(map(int, session["answered_questions"]))
    theta = float(session.get("theta", 0.0))

    rows = get_all_item_params()
    all_item_params = [tuple(r[1:]) for r in rows]
    answered_params = [all_item_params[i] for i in answered] if answered else []

    icc_path = plot_icc(answered_params, f"static/icc_{uuid.uuid4().hex}.png") if answered_params else None
    info_path = plot_item_information(answered_params, f"static/info_{uuid.uuid4().hex}.png") if answered_params else None

    n_total = len(responses)
    n_correct = sum(1 for r in responses if r == 1)
    n_wrong = n_total - n_correct
    accuracy = round((n_correct / n_total) * 100, 1) if n_total else 0

    if answered_params:
        se = theta_se(theta, answered_params)
        ci68 = (max(-4, theta - se), min(4, theta + se))
        ci95 = (max(-4, theta - 1.96 * se), min(4, theta + 1.96 * se))
    else:
        se, ci68, ci95 = None, None, None

    def ability_band(t):
        if t < -2: return "خیلی پایین"
        if t < -1: return "پایین"
        if t <= 1: return "متوسط"
        if t <= 2: return "بالا"
        return "خیلی بالا"

    role = session.get("role", "learner")
    pid = int(session["participant_id"])
    post_test_url = None
    post_test_done = False

    with sqlite3.connect(DATABASE, timeout=30) as conn:
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA busy_timeout=30000;")
        cur = conn.cursor()

        if role == "learner":
            post_test_url = url_for("post_test")
            post_test_done = cur.execute("SELECT 1 FROM strategy_answers WHERE participant_id=? LIMIT 1", (pid,)).fetchone() is not None
        elif role == "teacher":
            post_test_url = url_for("post_test_teacher")
            post_test_done = cur.execute("SELECT 1 FROM teacher_post_answers WHERE participant_id=? LIMIT 1", (pid,)).fetchone() is not None
        elif role == "manager":
            post_test_url = url_for("post_test_manager")
            post_test_done = cur.execute("SELECT 1 FROM manager_post_answers WHERE participant_id=? LIMIT 1", (pid,)).fetchone() is not None

    return render_template(
        "result.html",
        theta=theta,
        band=ability_band(theta),
        se=se,
        ci68=ci68,
        ci95=ci95,
        n_total=n_total,
        n_correct=n_correct,
        n_wrong=n_wrong,
        accuracy=accuracy,
        user_name=session.get("user_name", "کاربر ناشناس"),
        icc_image=icc_path,
        info_image=info_path,
        interpretation="—",
        has_post_test=post_test_done,
        role=role,
        post_test_url=post_test_url,
        post_test_done=post_test_done,
    )

# ----------------------------- پس‌آزمون مدرس -----------------------------
@app.route("/post_test_teacher", methods=["GET", "POST"])
def post_test_teacher():
    if "participant_id" not in session or session.get("role") != "teacher":
        return redirect(url_for("index"))

    pid = int(session["participant_id"])
    txt_col = pick_text_column("teacher_post_questions")

    with get_db_connection() as conn:
        cur = conn.cursor()

        if cur.execute("SELECT 1 FROM teacher_post_answers WHERE participant_id=? LIMIT 1", (pid,)).fetchone():
            return render_template("post_test_completed.html", role="teacher")

        questions = cur.execute(
            f"""
            SELECT id,
                   {txt_col} AS text,
                   COALESCE(question_type,'open') AS question_type,
                   COALESCE(is_required,1) AS is_required
            FROM teacher_post_questions
            WHERE COALESCE(is_active,1)=1
            ORDER BY COALESCE(display_order,0), id
            """
        ).fetchall()

        if request.method == "POST":
            errors = {}
            answers = {}

            for q in questions:
                key = f"q_{q['id']}"
                val = request.form.get(key)
                is_req = int(q["is_required"]) if q["is_required"] is not None else 1
                if is_req == 1 and (val is None or str(val).strip() == ""):
                    errors[int(q["id"])] = "این مورد الزامی است."
                else:
                    answers[int(q["id"])] = val

            if errors:
                return render_template("post_test_teacher.html", questions=questions, errors=errors, values=request.form)

            for qid, val in answers.items():
                q = next((qq for qq in questions if int(qq["id"]) == int(qid)), None)
                qtype = ((q["question_type"] if q else "open") or "open").lower()

                if qtype == "likert":
                    try:
                        aval = int(val) if val is not None and str(val).strip() != "" else None
                    except ValueError:
                        aval = None
                    atext = None
                else:
                    aval = None
                    atext = val

                cur.execute(
                    """
                    INSERT INTO teacher_post_answers (participant_id, question_id, answer_value, answer_text)
                    VALUES (?, ?, ?, ?)
                    ON CONFLICT(participant_id, question_id) DO UPDATE SET
                        answer_value=excluded.answer_value,
                        answer_text=excluded.answer_text,
                        created_at=CURRENT_TIMESTAMP
                    """,
                    (pid, int(qid), aval, atext),
                )

            conn.commit()
            return redirect(url_for("result"))

    return render_template("post_test_teacher.html", questions=questions, errors={}, values={})

# ----------------------------- پس‌آزمون مدیر -----------------------------
@app.route("/post_test_manager", methods=["GET", "POST"])
def post_test_manager():
    if "participant_id" not in session or session.get("role") != "manager":
        return redirect(url_for("index"))

    pid = int(session["participant_id"])
    txt_col = pick_text_column("manager_post_questions")

    with get_db_connection() as conn:
        cur = conn.cursor()

        if cur.execute("SELECT 1 FROM manager_post_answers WHERE participant_id=? LIMIT 1", (pid,)).fetchone():
            return render_template("post_test_completed.html", role="manager")

        questions = cur.execute(
            f"""
            SELECT id,
                   {txt_col} AS text,
                   COALESCE(question_type,'open') AS question_type,
                   COALESCE(is_required,1) AS is_required
            FROM manager_post_questions
            WHERE COALESCE(is_active,1)=1
            ORDER BY COALESCE(display_order,0), id
            """
        ).fetchall()

        if request.method == "POST":
            errors = {}
            answers = {}

            for q in questions:
                key = f"q_{q['id']}"
                val = request.form.get(key)
                is_req = int(q["is_required"]) if q["is_required"] is not None else 1
                if is_req == 1 and (val is None or str(val).strip() == ""):
                    errors[int(q["id"])] = "این مورد الزامی است."
                else:
                    answers[int(q["id"])] = val

            if errors:
                return render_template("post_test_manager.html", questions=questions, errors=errors, values=request.form)

            for qid, val in answers.items():
                q = next((qq for qq in questions if int(qq["id"]) == int(qid)), None)
                qtype = ((q["question_type"] if q else "open") or "open").lower()

                if qtype == "likert":
                    try:
                        aval = int(val) if val is not None and str(val).strip() != "" else None
                    except ValueError:
                        aval = None
                    atext = None
                else:
                    aval = None
                    atext = val

                cur.execute(
                    """
                    INSERT INTO manager_post_answers (participant_id, question_id, answer_value, answer_text)
                    VALUES (?, ?, ?, ?)
                    ON CONFLICT(participant_id, question_id) DO UPDATE SET
                        answer_value=excluded.answer_value,
                        answer_text=excluded.answer_text,
                        answered_at=CURRENT_TIMESTAMP
                    """,
                    (pid, int(qid), aval, atext),
                )

            conn.commit()
            return redirect(url_for("result"))

    return render_template("post_test_manager.html", questions=questions, errors={}, values={})

# ----------------------------- دانلود -----------------------------
@app.route("/download/<filetype>")
def download(filetype):
    if "responses" not in session or "answered_questions" not in session:
        return redirect(url_for("index"))

    responses = list(map(int, session["responses"]))
    answered = list(map(int, session["answered_questions"]))
    theta = float(session.get("theta", 0.0))

    filename = f"results_{uuid.uuid4().hex}"
    if filetype == "excel":
        filepath = f"static/{filename}.xlsx"
        save_results_to_excel(filepath, responses, answered, theta)
        return send_file(filepath, as_attachment=True)
    if filetype == "word":
        filepath = f"static/{filename}.docx"
        save_results_to_word(filepath, responses, answered, theta)
        return send_file(filepath, as_attachment=True)
    return redirect(url_for("result"))

# ----------------------------- اجرا -----------------------------
if __name__ == "__main__":
    init_db()
    app.run(debug=True, use_reloader=False, threaded=False)
