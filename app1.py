from flask import Flask, render_template, request, session, redirect, url_for
import os
import pandas as pd
from docx import Document
import matplotlib.pyplot as plt
import sqlite3

from flask import send_from_directory


app = Flask(__name__)
app.secret_key = 'your_secret_key'

OUTPUT_FOLDER = 'static/results'
if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

# ---------- توابع کمکی ----------

def save_results_to_excel(file_path, responses, item_params, theta):
    df = pd.DataFrame({
        'Item': range(1, len(responses) + 1),
        'Response': responses,
        'a': [ip[0] for ip in item_params] if item_params else [],
        'b': [ip[1] for ip in item_params] if item_params else [],
        'c': [ip[2] for ip in item_params] if item_params else [],
    })
    summary = pd.DataFrame({'Theta': [theta]})
    with pd.ExcelWriter(file_path) as writer:
        df.to_excel(writer, sheet_name='Responses', index=False)
        summary.to_excel(writer, sheet_name='Summary', index=False)

def save_results_to_word(file_path):
    doc = Document()
    doc.add_heading('نتایج آزمون انطباقی', 0)
    doc.add_paragraph('این فایل نتایج آزمون انطباقی را نشان می‌دهد.')
    doc.save(file_path)

def plot_icc(item_params, file_path):
    plt.figure(figsize=(6,4))
    theta_range = [i/10 for i in range(-40, 41)]
    for idx, (a, b, c) in enumerate(item_params):
        p_theta = [c + (1-c) / (1 + pow(2.718, -1.7*a*(t - b))) for t in theta_range]
        plt.plot(theta_range, p_theta, label=f'سؤال {idx+1}')
    plt.title('نمودار تابع احتمال پاسخ صحیح (ICC)')
    plt.xlabel('توانایی (θ)')
    plt.ylabel('احتمال پاسخ صحیح')
    plt.legend()
    plt.tight_layout()
    plt.savefig(file_path)
    plt.close()

def plot_item_information(item_params, file_path):
    plt.figure(figsize=(6,4))
    theta_range = [i/10 for i in range(-40, 41)]
    for idx, (a, b, c) in enumerate(item_params):
        info = []
        for t in theta_range:
            p = c + (1-c) / (1 + pow(2.718, -1.7*a*(t - b)))
            info_val = (1.7**2) * (a**2) * p * (1-p) / ((1-c)**2) if p else 0
            info.append(info_val)
        plt.plot(theta_range, info, label=f'سؤال {idx+1}')
    plt.title('نمودار اطلاعات آیتم')
    plt.xlabel('توانایی (θ)')
    plt.ylabel('اطلاعات آیتم')
    plt.legend()
    plt.tight_layout()
    plt.savefig(file_path)
    plt.close()

def insert_user_result(full_name, native_language, major, age, persian_familiarity, theta, responses):
    conn = sqlite3.connect('results.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS results (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        full_name TEXT,
        native_language TEXT,
        major TEXT,
        age INTEGER,
        persian_familiarity TEXT,
        theta REAL,
        responses TEXT
    )''')
    responses_str = ','.join(map(str, responses))
    c.execute('INSERT INTO results (full_name, native_language, major, age, persian_familiarity, theta, responses) VALUES (?, ?, ?, ?, ?, ?, ?)',
              (full_name, native_language, major, age, persian_familiarity, theta, responses_str))
    conn.commit()
    conn.close()

# ---------- روت‌ها ----------

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/test', methods=['POST'])
def start_test():
    data = request.form
    session['full_name'] = data.get('full_name')
    session['native_language'] = data.get('native_language')
    session['major'] = data.get('major')
    session['age'] = data.get('age')
    session['persian_familiarity'] = data.get('persian_familiarity')

    session['responses'] = []
    session['item_params'] = [(1.0, 0.0, 0.2)] * 30  # نمونه پارامتر سوال‌ها (a,b,c)
    session['theta'] = 0

    return redirect(url_for('test_question'))

@app.route('/test_question', methods=['GET', 'POST'])
def test_question():
    if 'responses' not in session:
        return redirect(url_for('index'))

    if request.method == 'POST':
        answer = request.form.get('answer')
        session['responses'].append(answer)

        # به‌روزرسانی فرضی theta
        session['theta'] += 0.1

        if len(session['responses']) >= 30:
            return redirect(url_for('result'))

        return redirect(url_for('test_question'))

    question_number = len(session['responses']) + 1
    question = {
        'id': question_number,
        'text': f'متن سؤال شماره {question_number} اینجا قرار می‌گیرد.',
        'options': ['گزینه ۱', 'گزینه ۲', 'گزینه ۳', 'گزینه ۴']
    }
    progress_percentage = int((len(session['responses']) / 30) * 100)

    return render_template('test_question.html',
                           question=question,
                           question_number=question_number,
                           progress_percentage=progress_percentage)

@app.route('/download/<filetype>/<filename>')
def download_file(filetype, filename):
    # آدرس فولدر فایل‌ها
    directory = os.path.join(app.root_path, 'static', 'results')
    
    # فقط اجازه پسوندهای امن را می‌دهیم (مثلاً xlsx و docx)
    allowed_types = ['excel', 'word']
    allowed_extensions = {'excel': '.xlsx', 'word': '.docx'}
    
    if filetype not in allowed_types:
        return "نوع فایل نامعتبر است.", 400
    
    if not filename.endswith(allowed_extensions[filetype]):
        return "پسوند فایل نامعتبر است.", 400

    try:
        return send_from_directory(directory, filename, as_attachment=True)
    except FileNotFoundError:
        return "فایل پیدا نشد.", 404



@app.route('/result')
def result():
    if 'responses' not in session:
        return redirect(url_for('index'))

    theta = session.get('theta', 0)
    responses = session.get('responses', [])
    item_params = session.get('item_params', [])
    full_name = session.get('full_name', 'نامشخص')

    name = full_name.replace(' ', '_')
    excel_file = f'result_{name}.xlsx'
    word_file = f'result_{name}.docx'
    icc_file = f'icc_{name}.png'
    info_file = f'info_{name}.png'

    save_results_to_excel(os.path.join(OUTPUT_FOLDER, excel_file), responses, item_params, theta)
    save_results_to_word(os.path.join(OUTPUT_FOLDER, word_file))
    plot_icc(item_params, os.path.join(OUTPUT_FOLDER, icc_file))
    plot_item_information(item_params, os.path.join(OUTPUT_FOLDER, info_file))

    # ذخیره در دیتابیس
    insert_user_result(
        full_name=full_name,
        native_language=session.get('native_language'),
        major=session.get('major'),
        age=int(session.get('age')) if session.get('age') else None,
        persian_familiarity=session.get('persian_familiarity'),
        theta=theta,
        responses=responses
    )

    return render_template('result.html',
                           full_name=full_name,
                           native_language=session.get('native_language'),
                           major=session.get('major'),
                           age=session.get('age'),
                           persian_familiarity=session.get('persian_familiarity'),
                           theta=theta,
                           excel_file=url_for('static', filename=f'results/{excel_file}'),
                           word_file=url_for('static', filename=f'results/{word_file}'),
                           icc_image=url_for('static', filename=f'results/{icc_file}'),
                           info_image=url_for('static', filename=f'results/{info_file}'))

if __name__ == '__main__':
    app.run(debug=True)
